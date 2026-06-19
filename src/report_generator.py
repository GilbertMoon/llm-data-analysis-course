from pathlib import Path

from docx import Document


def create_word_report(
    title: str,
    summary: str,
    insights: list[str],
    output_path: str | Path = "reports/outputs/analysis_report.docx",
) -> Path:
    """분석 제목, 요약, 인사이트 목록으로 간단한 Word 보고서를 생성합니다."""
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    document.add_heading(title, level=1)
    document.add_heading("분석 요약", level=2)
    document.add_paragraph(summary)
    document.add_heading("주요 인사이트", level=2)

    for insight in insights:
        document.add_paragraph(insight, style="List Bullet")

    document.save(path)
    return path
